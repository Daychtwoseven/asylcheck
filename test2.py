from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, Length, Cm

# Open a new document
doc = Document()

# Add a table with one row and two cells
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False

# Get the first row of the table
row = table.rows[0]

# Add content to the first column
left_cell = row.cells[0]
left_paragraph = left_cell.paragraphs[0]
left_paragraph.add_run("Bundesamt für Fremdenwesen und Asyl - Direktion \nModecenterstraße 22 \n1030 Wien \n \n")

left_paragraph.add_run("per E-Mail: BFA-Einlaufstelle@bmi.gv.at").bold = True

left_paragraph.add_run("\n \nGebühreneinzug \nAT63 2081 5232 0008 4582\nSTSPAT2GXXX \n").font.size = Pt(10)

# Add the first word
run1 = left_paragraph.add_run("\n[Datum]")
run1.alignment = 0  # Set alignment if desired

# Add space between the words
space = ' ' * 90  # Adjust the number of spaces as needed
run_space = left_paragraph.add_run(space)
run_space.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  #

# Add the second word
run2 = left_paragraph.add_run("GZ IPZ: [Kartennummer]\n \n")
run2.alignment = 3  # Set alignment if desired

# Add the first word
run1 = left_paragraph.add_run("Beschwerdeführer:")
run1.underline = True
run1.alignment = 0  # Set alignment if desired

# Add the second word
run2 = left_paragraph.add_run("    [Vorname, Nachname, Geburtsdatum]")
run2.alignment = 3  # Set alignment if desired

left_cell = row.cells[0]
s = left_cell.add_paragraph("")
s.add_run("vertreten durch:").underline = True
s.paragraph_format.left_indent = Inches(1.4)


left_cell = row.cells[0]
s = left_cell.add_paragraph("")
s.add_run("Held Berdnik Astner & Partner\nRechtsanwälte GmbH\nAT-1090 Wien, Rooseveltplatz 10\nCode "
          "P034285\n\nVollmacht gemäß §§ 8, 21e RAO\nund § 30 Abs 2 ZPO erteilt\nKosten gemäß § 19a RAO\nzu Handen "
          "der Vertreterin").font.size = Pt(8)
s.paragraph_format.left_indent = Inches(1.4)

# --------------------------------------------------------------------------------------------------------------------
s = left_cell.add_paragraph("")
s.add_run("Säumige Behörde:").underline = True
s.underline = True

# Add the second word
s.add_run("    Bundesamt für Fremdenwesen und Asyl - Direktion \n                                       Modecenterstraße 22, A-1030 Wien\n\n\n")

# --------------------------------------------------------------------------------------------------------------------

# --------------------------------------------------------------------------------------------------------------------
s = left_cell.add_paragraph("")
s.add_run("wegen:").underline = True
s.underline = True

# Add the second word
s.add_run("                         Antrag auf internationalen Schutz - Säumnis\n\n")


s = left_cell.add_paragraph("")
s.add_run("                                       I.      VOLLMACHTSBEKANNTGABE").bold = True

s = left_cell.add_paragraph("")
s.add_run("                                            II.     SÄUMNISBESCHWERDE").bold = True
# --------------------------------------------------------------------------------------------------------------------

# Add content to the second column
right_cell = row.cells[1]
right_paragraph = right_cell.paragraphs[0]
right_paragraph.add_run("Right column content")

for cell in table.columns[0].cells:
    cell.width = Inches(5)


section = doc.sections[0]

# Access the footer of the section
footer = section.footer

# Add a paragraph to the footer
paragraph = footer.paragraphs[0]
run = paragraph.add_run("Seite 1 | 4")
run.font.size = Pt(10)  # Set the font size
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Set the alignment


# Save the document
doc.save('path_to_save_document.docx')
