from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt


def generate_docx():
    try:
        document = Document()

        # Add a table with one row and two cells
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False

        # Get the first row of the table
        row = table.rows[0]

        # Add content to the first column
        left_cell = row.cells[0]
        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.add_run(
            "Bundesamt für Fremdenwesen und Asyl - Direktion \nModecenterstraße 22 \n1030 Wien \n \n")

        left_paragraph.add_run("per E-Mail: BFA-Einlaufstelle@bmi.gv.at").bold = True

        left_paragraph.add_run("\n \nGebühreneinzug \nAT63 2081 5232 0008 4582\nSTSPAT2GXXX \n").font.size = Pt(10)

        # Add the first word
        run1 = left_paragraph.add_run("\n[Datum]")
        run1.alignment = 0  # Set alignment if desired

        # Add space between the words
        space = ' ' * 75  # Adjust the number of spaces as needed
        run_space = left_paragraph.add_run(space)
        run_space.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  #

        # Add the second word
        run2 = left_paragraph.add_run("GZ IPZ: [Kartennummer]\n \n")
        run2.alignment = 3  # Set alignment if desired

        # Add the first word
        run1 = left_paragraph.add_run("Beschwerdeführer:")
        run1.underline = True
        run1.alignment = 0  # Set alignment if desired

        # Add the second word
        run2 = left_paragraph.add_run("    [Vorname, Nachname, Geburtsdatum]")
        run2.alignment = 3  # Set alignment if desired

        left_cell = row.cells[0]
        s = left_cell.add_paragraph("")
        s.add_run("vertreten durch:").underline = True
        s.paragraph_format.left_indent = Inches(1.4)

        left_cell = row.cells[0]
        s = left_cell.add_paragraph("")
        s.add_run().add_picture('hba.jpg', width=Pt(23), height=Pt(10))
        s.add_run("\nHeld Berdnik Astner & Partner\nRechtsanwälte GmbH\nAT-1090 Wien, Rooseveltplatz 10\nCode "
                  "P034285\n\nVollmacht gemäß §§ 8, 21e RAO\nund § 30 Abs 2 ZPO erteilt\nKosten gemäß § 19a RAO\nzu Handen "
                  "der Vertreterin").font.size = Pt(8)
        s.paragraph_format.left_indent = Inches(1.4)

        # --------------------------------------------------------------------------------------------------------------------
        s = left_cell.add_paragraph("")
        s.add_run("Säumige Behörde:").underline = True
        s.underline = True

        # Add the second word
        s.add_run(
            "      Bundesamt für Fremdenwesen und Asyl - Direktion \n                                         Modecenterstraße "
            "22, A-1030 Wien\n\n\n")

        # --------------------------------------------------------------------------------------------------------------------

        # --------------------------------------------------------------------------------------------------------------------
        s = left_cell.add_paragraph("")
        s.add_run("wegen:").underline = True
        s.underline = True

        # Add the second word
        s.add_run("                           Antrag auf internationalen Schutz - Säumnis\n\n")

        s = left_cell.add_paragraph("")
        s.add_run("                                         I.      VOLLMACHTSBEKANNTGABE").bold = True

        s = left_cell.add_paragraph("")
        s.add_run("                                              II.     SÄUMNISBESCHWERDE").bold = True
        # --------------------------------------------------------------------------------------------------------------------

        # Add content to the second column
        right_cell = row.cells[1]
        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.add_run().add_picture('hba.jpg', width=Pt(35), height=Pt(15))
        right_paragraph.add_run(
            "\nHeld Berdnik Astner & Partner\nRechtsanwälte GmbH\n\nDr. Guido HELD em. RA\nRA Mag. Lukas "
            "HELD, LL.M.\nRA Dr. Gottfried BERDNIK\nRA Dr. Bernhard ASTNER\nRA Dr. Joachim ZIERLER\n[auch "
            "als Steuerberater zugelassen]\nRA Dr. Ullrich SAURER\nRA Mag. Johannes ZINK\nRA Dr. Leo "
            "GRÖTSCHNIG\nRA Dr. Robert MIKLAUSCHINA\n[auch gerichtlicher Sachverständiger für\nLogistik-, "
            "Speditions-, Frachtrecht]\nRA Mag. Dieter HUTTER\nRA Dr. Thomas GRUBER, LL.M.\nRA Dr. "
            "Philipp SPATZ, LL.M.\n[auch in New York zugelassen]\nRA Mag. Michael WOHLGEMUTH, "
            "LL.M.\nRA Dr. Peter IVANKOVICS, M.A., LL.M.\nRA Mag. Lukas HONZAK\nRA Mag. Philipp "
            "WIESER\nRA Mag. David JODLBAUER\nRA Mag. Barbara KAILBAUER\nRA Mag. Michael STEINER\nRA Mag. "
            "Michaela STÜCKLER\nRA Dr. Isabella FANK\n\nUniv.-Prof. Dr. Gert-Peter REISSNER\n["
            "Konsulent]\n\nAT–1090 WIEN\nRooseveltplatz 10\nT +43 (0) 50 8060 400\nE "
            "vienna@hba.at\n\nAT–8010 "
            "GRAZ\nKarmeliterplatz 4\nT +43 (0) 50 8060 200\nE graz@hba.at\n\nAT–9020 "
            "KLAGENFURT\nTheaterplatz 5\nT +43 (0) 50 8060 600\nE klagenfurt@hba.at\n\nAT–7000 "
            "EISENSTADT\nMarktstraße 2\nT +43 (0) 50 8060 700\nE eisenstadt@hba.at\n\nwww.hba.at\n\nFN "
            "253765i | Landesgericht fü\nZivilrechtssachen Graz | Sitz "
            "Graz\n\nSammelanderkonto\nSteiermärkische Bank und\nSparkassen AG\nIBAN AT23 2081 5232 0008 "
            "1034\nBIC STSPAT2GXXX\n\nATU 61172028 ")
        for run in right_paragraph.runs:
            run.font.size = Pt(7)
        for cell in table.columns[0].cells:
            cell.width = Inches(5)
        # --------------------------------------------------------------------------------------------------------------------

        p = document.add_paragraph("\n\n")
        p.add_run("I.").bold = True
        p.alignment = 1
        document.add_paragraph(
            "In außen näher bezeichneter Angelegenheit gibt der Beschwerdeführer bekannt, die Held Berdnik "
            "Astner & Partner Rechtsanwälte GmbH, Rooseveltplatz 10, 1090 Wien, mit seiner "
            "rechtsfreundlichen Vertretung im Säumnisbeschwerdeverfahren beauftragt und bevollmächtigt zu "
            "haben. Die außen angeführte Rechtsanwaltskanzlei beruft sich gemäß § 8 Abs 1 RAO iVm § 10 AVG "
            "auf die ihr erteilte Vollmacht.")

        p = document.add_paragraph("")
        p.add_run("II.").bold = True
        p.size = Pt(4)
        p.alignment = 1

        document.add_paragraph("Infolge des Ablaufs der Entscheidungsfrist erhebt der Beschwerdeführer nachstehende")
        p = document.add_paragraph("")
        p.add_run("SÄUMNISBESCHWERDE").bold = True
        p.alignment = 1
        for run in p.runs:
            run.font.size = Pt(14)

        document.add_paragraph("an das Bundesverwaltungsgericht und führt hierzu an:")

        # Add the numbered list
        numbered_list = [1, 2, 3, 4, 5]  # Example list of numbers
        for number in numbered_list:
            if number == 1:
                list_paragraph = document.add_paragraph()
                list_paragraph.style = 'List Number'
                list_run = list_paragraph.add_run("Sachverhalt").bold = True
                p = document.add_paragraph("Am [Datum der Erstbefragung bzw Datum der Ausweisausstellung] hat der "
                                           "Beschwerdeführer seine Erstbefragung nach dem AsylG absolviert, spätestens mit "
                                           "diesem Zeitpunkt hat er daher eine Entscheidung über seinen Antrag auf "
                                           "Zuerkennung internationalen Schutzes nach dem AsylG beantragt.")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "Die Aufenthaltsberechtigungskarte gem § 51 AsylG wurde dem Beschwerdeführer ausgestellt, in den letzten "
                    "[Zeitraum zwischen Säumnisbeschwerde und Erstbefragung] Monaten wurde aber keine Entscheidung getroffen. "
                    "Auch andere behördliche Tätigkeit wodurch die massiven Verzögerungen zu erklären sind, waren nicht "
                    "erkennbar für den Beschwerdeführer. ")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph("Beweis:	beizuschaffender BFA-Akt")
            elif number == 2:
                list_paragraph = document.add_paragraph()
                list_paragraph.style = 'List Number'
                list_run = list_paragraph.add_run("Rechtliche Überlegungen").bold = True

                p = document.add_paragraph(
                    "Gemäß Art. 130 Abs. 1 Z 3 B-VG erkennen die Verwaltungsgerichte über Beschwerden "
                    "wegen Verletzung der Entscheidungspflicht durch eine Verwaltungsbehörde.")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "Gemäß Art. 131 Abs. 2 B-VG erkennt das BVwG über Beschwerden gemäß Art. 130 Abs. 1 in Rechtssachen in "
                    "den Angelegenheiten der Vollziehung des Bundes, die unmittelbar von Bundesbehörden besorgt werden.")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "GGemäß § 7 Abs. 1 Z 4 BFA-VG entscheidet das Bundesverwaltungsgericht über Beschwerden wegen Verletzung "
                    "der Entscheidungspflicht des Bundesamtes für Fremdenwesen und Asyl.")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "Gemäß § 8 Abs. 1 VwGVG kann Beschwerde wegen Verletzung der Entscheidungspflicht gemäß Art. 130 Abs. 1 Z "
                    "3 B-VG (Säumnisbeschwerde) erst erhoben werden, wenn die Behörde die Sache nicht innerhalb von sechs "
                    "Monaten, wenn gesetzlich eine kürzere oder längere Entscheidungsfrist vorgesehen ist, innerhalb dieser "
                    "entschieden hat. Die Frist beginnt mit dem Zeitpunkt, in dem der Antrag auf Sachentscheidung bei der "
                    "Stelle eingelangt ist, bei der er einzubringen war. Die Beschwerde ist abzuweisen, wenn die Verzögerung "
                    "nicht auf ein überwiegendes Verschulden der Behörde zurückzuführen ist.")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "Gemäß § 9 Abs. 1 VwGVG hat eine Beschwerde zu enthalten:")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "1. die Bezeichnung des angefochtenen Bescheides, der angefochtenen Ausübung unmittelbarer "
                    "verwaltungsbehördlicher Befehls- und Zwangsgewalt oder der angefochtenen Weisung,")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "2. die Bezeichnung der belangten Behörde,")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "3. die Gründe, auf die sich die Behauptung der Rechtswidrigkeit stützt,")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "4. das Begehren und")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "5. die Angaben, die erforderlich sind, um zu beurteilen, ob die Beschwerde rechtzeitig eingebracht ist.")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "Gemäß Abs. 2 Z. 3 ist belangte Behörde in den Fällen des Art. 130 Abs. 1 Z 3 B-VG jene Behörde, "
                    "die den Bescheid nicht erlassen hat.")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph(
                    "Gemäß Abs. 5 entfallen bei Beschwerden wegen Verletzung der Entscheidungspflicht gemäß Art. 130 Abs. 1 Z "
                    "3 B-VG die Angaben nach Abs. 1 Z 1 bis 3 und 5. Als belangte Behörde ist die Behörde zu bezeichnen, "
                    "deren Entscheidung in der Rechtssache begehrt wurde. Ferner ist glaubhaft zu machen, dass die Frist zur "
                    "Erhebung der Säumnisbeschwerde gemäß § 8 Abs. 1 abgelaufen ist.")
                p.paragraph_format.left_indent = Inches(0.5)

            elif number == 3:
                list_paragraph = document.add_paragraph()
                list_paragraph.style = 'List Number'
                list_run = list_paragraph.add_run("Antrag").bold = True

                p = document.add_paragraph(
                    "Der Beschwerdeführer stellt daher den")
                p.paragraph_format.left_indent = Inches(0.5)

                p = document.add_paragraph("")
                p.add_run("ANTRAG").bold = True
                p.alignment = 1
                for run in p.runs:
                    run.font.size = Pt(14)

                p = document.add_paragraph(
                    "Das Bundesverwaltungsgericht möge,")
                p.paragraph_format.left_indent = Inches(0.5)

                list_items = [1, 2]  # Example list items
                # Add the list items
                for item in list_items:
                    if item == 1:
                        inside_list_paragraph = document.add_paragraph()
                        inside_list_paragraph.paragraph_format.left_indent = Inches(1.3)
                        inside_list_paragraph.style = 'List Bullet'
                        list_run = inside_list_paragraph.add_run("der Säumnisbeschwerde stattgeben, die Verletzung der "
                                                                 "Entscheidungspflicht von Seiten des Bundesamts für "
                                                                 "Fremdenwesen und Asyl feststellen und gemäß § 28 Abs 7 1. "
                                                                 "Fall VwGVG in der Sache selbst zu entscheiden; in eventu. \n")

                    elif item == 2:
                        inside_list_paragraph = document.add_paragraph()
                        inside_list_paragraph.paragraph_format.left_indent = Inches(1.3)
                        inside_list_paragraph.style = 'List Bullet'
                        list_run = inside_list_paragraph.add_run("der Säumnisbeschwerde stattgeben, die Verletzung der "
                                                                 "Entscheidungspflicht von Seiten des Bundesamts für "
                                                                 "Fremdenwesen und Asyl feststellen und der säumigen Behörde "
                                                                 "gemäß § 28 Abs 7 2. Fall VwGVG aufzutragen, den versäumten "
                                                                 "Bescheid binnen 8 Wochen zu erlassen.")

        p = document.add_paragraph("[Vorname Nachname]")
        p.alignment = 2

        # Specify the desired paper size
        paper_width = 8.5  # Width in inches
        paper_height = 11  # Height in inches

        # Access the section properties
        for section in document.sections:
            # Change the paper size
            section.page_width = Pt(paper_width * 72)
            section.page_height = Pt(paper_height * 72)

        # Iterate through sections
        for section in document.sections:
            # Access the footer of the section
            footer = section.footer

            # Add a paragraph to the footer
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Set the alignment

            # Create a field code for the page number
            fld_page = OxmlElement('w:fldSimple')
            fld_page.set(qn('w:instr'), 'PAGE')

            # Create a run for the field code
            run_page = paragraph.add_run()
            run_page._r.append(fld_page)

        document.save("test.docx")
    except Exception as e:
        print(e)

generate_docx()